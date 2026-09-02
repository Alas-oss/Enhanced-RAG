from __future__ import annotations

from pathlib import Path
TEXT_EXTENSIONS = {".txt", ".md"}

def load_document(path: str) -> str:
    p = Path(path)
    ext = p.suffix.lower()

    if ext in TEXT_EXTENSIONS:
        return p.read_text(encoding="utf-8", errors="ignore")
    if ext == ".pdf":
        return _load_pdf(p)
    if ext == ".docx":
        return _load_docx(p)

    return p.read_text(encoding="utf-8", errors="ignore")

def _load_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ImportError(
            "Reading PDF files required pypdf. Install with: pip install pypdf"
        ) from exc

    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        pages.append(f"\n\n[Page {i + 1}]\n{text}")
    return "".join(pages).strip()

def _load_docs(path: Path) -> str:
    try: 
        import docx
    except ImportError as exc:
        raise ImportError(
            "Reading DOCX files requires python-docx. Install with: pip install python-docx"
        ) from exc

    document = docx.Document(str(path))
    parts = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if rows: 
            parts.append("\n".join(rows))

    return "\n\n".join(parts)
